"""
Скрипт для извлечения текста из .docx файлов в raw/
Запускать из командной строки: python extract_docx.py
"""
import sys

try:
    from docx import Document
except ImportError:
    print("Устанавливаю python-docx...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document

import os

RAW_DIR = r"C:\Users\user\Desktop\MyBusiness\raw"
OUTPUT_DIR = r"C:\Users\user\Desktop\MyBusiness\raw"

files = [
    "2026-03-01_стратегический-отчет-дорожная-карта.docx",
    "2026-03-01_анализ-рынка-состояние-и-перспективы.docx",
    "2026-03-01_анализ-незанятых-ниш.docx",
]

for fname in files:
    path = os.path.join(RAW_DIR, fname)
    if not os.path.exists(path):
        print(f"НЕ НАЙДЕН: {fname}")
        continue

    print(f"\n{'='*70}")
    print(f"ФАЙЛ: {fname}")
    print('='*70)

    doc = Document(path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    text = "\n".join(paragraphs)
    print(text[:5000])
    
    if len(text) > 5000:
        print(f"\n... [ещё {len(text)-5000} символов — скопируй вторую половину если нужно]")
    
    # Сохраняем .txt рядом с .docx
    out_path = os.path.join(OUTPUT_DIR, fname.replace(".docx", "_extracted.txt"))
    with open(out_path, "w", encoding="utf-8") as f:
        f.write(text)
    print(f"\n✅ Сохранён: {out_path}")

print("\n\nГОТОВО. Вставь весь вывод выше в чат с Claude.")
